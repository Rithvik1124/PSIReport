from fastapi import FastAPI
from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware

app = FastAPI()

# Add CORS *before* any routes
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
from fastapi.responses import FileResponse
from pydantic import BaseModel
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from docx import Document
import re
import datetime
from docx.enum.text import WD_COLOR_INDEX
import zipfile
from PIL import Image
import io
from openai import OpenAI
import time
import base64
import os



client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))


class URLRequest(BaseModel):
    url: str

def setup_driver(mobile=False):
    chrome_options = Options()
    chrome_options.add_argument("--headless=new")  # or "--headless"
    chrome_options.add_argument("--window-size=1000,800")
    chrome_options.add_argument("--disable-gpu")
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--disable-dev-shm-usage")  # 👈 Railway-specific
    chrome_options.binary_location = "/usr/bin/chromium"

    if mobile:
        chrome_options.add_experimental_option("mobileEmulation", {"deviceName": "Pixel 2"})        

    service = Service("/usr/bin/chromedriver")
    return webdriver.Chrome(service=service, options=chrome_options)

def extract_data(driver):
    selectors = {
        'Performance Score': '#performance .lh-exp-gauge__percentage',
        'Accessibility':'#accessibility .lh-gauge__percentage',
        'SEO':'#seo .lh-gauge__percentage',
        'BP':'#best-practices .lh-gauge__percentage',
        'LCP': '#largest-contentful-paint',
        'CLS': '#cumulative-layout-shift',
        'SI': '#speed-index',
        'TBT': '#total-blocking-time',
        'FCP': '#first-contentful-paint',
        'Diagnostics': '.lh-audit-group--diagnostics',
        'Insights': '.lh-audit-group--insights'
    }
    data = {}
    for label, selector in selectors.items():
        try:
            WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.CSS_SELECTOR, selector)))
            el = driver.find_element(By.CSS_SELECTOR, selector)
            driver.execute_script("arguments[0].scrollIntoView();", el)
            time.sleep(0.5)
            data[label] = el.text
        except Exception as e:
            data[label] = f"Error: {str(e)}"
            print(f"⚠️ Could not extract {label}: {e}")
    return data

def screenshot_to_pdf_base64(driver, path):
    # Scroll back to top of the page
    driver.execute_script("window.scrollTo(0, 0);")
    time.sleep(0.5)

    # Measure full page size
    total_width = driver.execute_script("return document.body.scrollWidth")
    total_height = driver.execute_script("return document.body.scrollHeight")

    # Resize browser to fit entire page
    driver.set_window_size(total_width, total_height)
    time.sleep(1)  # allow time for layout/render

    # Capture screenshot
    png_data = driver.get_screenshot_as_png()
    image = Image.open(io.BytesIO(png_data)).convert("RGB")
    image.save(path, "PDF")

def generate_advice(url, data_mob,data_desk):
    prompt = f"""
Act like an expert web performance consultant. I am running a Shopify website and I am a novice developer
who is not experienced with SEO or performance optimization. I will paste the output from a PageSpeed Insights report.
Go through the results and give me a structured improvement plan:
Audit Data (Mobile):
- Performance Score: {data_mob.get('Performance Score')}
- Accessibility: {data_mob.get('Accessibility')}
- Best Practices: {data_mob.get('BP')}
- SEO: {data_mob.get('SEO')}
- LCP: {data_mob.get('LCP')}
- CLS: {data_mob.get('CLS')}
- Speed Index: {data_mob.get('SI')}
- TBT: {data_mob.get('TBT')}
- FCP: {data_mob.get('FCP')}
- Diagnostics: {data_mob.get('Diagnostics')}
- Insights: {data_mob.get('Insights')}

Audit Data (Desktop):
- Performance Score: {data_desk.get('Performance Score')}
- Accessibility: {data_desk.get('Accessibility')}
- Best Practices: {data_desk.get('BP')}
- SEO: {data_desk.get('SEO')}
- LCP: {data_desk.get('LCP')}
- CLS: {data_desk.get('CLS')}
- Speed Index: {data_desk.get('SI')}
- TBT: {data_desk.get('TBT')}
- FCP: {data_desk.get('FCP')}
- Diagnostics: {data_desk.get('Diagnostics')}
- Insights: {data_desk.get('Insights')}
Break it down by category: Performance, Accessibility, Best Practices, SEO. Explain what the issue is in plain English.
Tell me why it matters and exactly how to fix it. Include examples of the code changes, server headers, or settings to adjust.
Don't assume I know SEO or Core Web Vitals—spell out every step.

URL: {url}

Give a detailed, plain-English optimization plan for Performance, Accessibility, SEO, and Best Practices.
Explain what’s wrong and exactly how to fix it.
"""

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a helpful web performance optimization expert."},
                {"role": "user", "content": prompt}
            ]
        )
        return response.choices[0].message.content

    except Exception as e:
        print("❌ OpenAI error:", e)
        return "Error generating advice from OpenAI."
def get_name(url):
    name = url.split('https://')[1].split('.')
    name = name[0] + str(datetime.date.today())
    return name

def parse_markdown_with_code(doc: Document, markdown: str):
    code_blocks = []
    def replacer(match):
        code_blocks.append(match.group(1).strip())
        return f"[[CODE_BLOCK_{len(code_blocks) - 1}]]"

    markdown = re.sub(r"```(?:\w*\n)?(.*?)```", replacer, markdown, flags=re.DOTALL)

    for line in markdown.splitlines():
        line = line.strip()
        if not line:
            continue

        if line.startswith("[[CODE_BLOCK_"):
            idx = int(re.search(r"\[\[CODE_BLOCK_(\d+)\]\]", line).group(1))
            add_code_block(doc, code_blocks[idx])
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("## ") or line.startswith("#### "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("- "):
            add_formatted_paragraph(doc, line[2:], style='List Bullet')
        else:
            add_formatted_paragraph(doc, line)

def add_formatted_paragraph(doc, text, style=None):
    paragraph = doc.add_paragraph(style=style)

    if text.startswith("<img"):
        run = paragraph.add_run(text)
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        return

    pattern = re.compile(r"\*\*(.*?)\*\*")
    pos = 0
    for match in pattern.finditer(text):
        before = text[pos:match.start()]
        bold_text = match.group(1)

        if before:
            paragraph.add_run(before)
        bold_run = paragraph.add_run(bold_text)
        bold_run.bold = True
        pos = match.end()

    if pos < len(text):
        paragraph.add_run(text[pos:])
    

def add_code_block(doc, block_text):
    for line in block_text.splitlines():
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.highlight_color = WD_COLOR_INDEX.TURQUOISE


@app.post("/analyze")
async def analyze(request: URLRequest):
    print("🌐 Received URL:", request.url)
    full_url = f"https://pagespeed.web.dev/analysis?url={request.url}"
    desktop_driver = None
    mobile_driver = None
    get_url_driver = None
    try:
        pdf_desktop_path="/tmp/screenshot_desktop.pdf"
        pdf_mobile_path="/tmp/screenshot_mobile.pdf"
        print("🚀 Loading PSI...")
        get_url_driver=setup_driver(mobile=False)
        get_url_driver.get(full_url)
        time.sleep(15)
        
        
        curr_url=get_url_driver.current_url
        desktop_driver = setup_driver(mobile=False)
        desktop_driver.get(curr_url)
        time.sleep(30)
        mobile_driver = setup_driver(mobile=True)
        mobile_driver.get(curr_url)
        time.sleep(30) 

        print("📊 Extracting metrics...")
        metrics_desk = extract_data(desktop_driver)
        metrics_mob = extract_data(mobile_driver)
        screenshot_to_pdf_base64(desktop_driver,pdf_desktop_path)
        screenshot_to_pdf_base64(mobile_driver,pdf_mobile_path)
        
        
        

        print("🤖 Getting AI advice...")
        advice = generate_advice(request.url, metrics_mob,metrics_desk)
        print("📝 AI advice content (first 300 chars):")
        print(advice[:300] if advice else "No advice returned")

        # 📄 Generate .docx
        # 📄 Generate docx with formatted AI advice only
        doc = Document()
        parse_markdown_with_code(doc, advice)
        filename1 = get_name(request.url) + ".docx"
        doc_path = f"/tmp/{filename1}"
        doc.save(doc_path)

        print("✅ Returning docx file")
        zip_path = "/tmp/psi_report_bundle.zip"
        with zipfile.ZipFile(zip_path, "w") as zipf:
            zipf.write(doc_path, arcname="psi_advice.docx")
            zipf.write(pdf_desktop_path, arcname="screenshot_desktop.pdf")
            zipf.write(pdf_mobile_path, arcname="screenshot_mobile.pdf")
        print("📦 Verifying ZIP file contents:")
        with zipfile.ZipFile(zip_path, "r") as zipf:
            print("Included files:", zipf.namelist())
        return FileResponse(path=zip_path, filename=f"psi_report_bundle.zip", media_type="application/zip")

    except Exception as e:
        print("🔥 CRITICAL ERROR:", e)
        return {"error": str(e)}

    finally:
        if desktop_driver:
            try:
                desktop_driver.quit()
            except Exception as e:
                print("⚠️ Error quitting desktop driver:", e)
        if mobile_driver:
            try:
                mobile_driver.quit()
            except Exception as e:
                print("⚠️ Error quitting mobile driver:", e)
        if get_url_driver:
            try:
                get_url_driver.quit()
            except Exception as e:
                print("⚠️ Error quitting mobile driver:", e)
@app.get('/')
def read_root():
    return {"message": "API is live"}
@app.get("/test")
def ping():
    return {"ping": "pong"}
